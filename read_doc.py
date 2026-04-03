import codecs
from docx import Document

doc = Document(r'e:\HR\expense_project\Expense Tracker Project – Modules.docx')

with codecs.open(r'e:\HR\expense_project\doc_content.txt', 'w', 'utf-8') as f:
    for p in doc.paragraphs:
        if p.text.strip():
            style = p.style.name if p.style else 'Normal'
            f.write(f"[{style}] {p.text}\n")
    
    for ti, table in enumerate(doc.tables):
        f.write(f"\n--- TABLE {ti+1} ---\n")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            f.write(" | ".join(cells) + "\n")
